"""
HackRx Document Intelligence System - Complete Cross-Platform Implementation
Matches exact problem statement requirements and API specification
"""

import requests
import PyPDF2
from docx import Document
import email
from bs4 import BeautifulSoup
import os
import logging
from typing import Tuple, List, Dict, Any, Optional
import re
import time
import hashlib
import json
from dotenv import load_dotenv
import google.generativeai as genai
from pinecone import Pinecone
from sentence_transformers import SentenceTransformer
import platform

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class HackRxDocumentIntelligenceSystem:
    """Complete Document Intelligence System matching HackRx specifications"""
    
    def __init__(self):
        """Initialize the system with cross-platform compatibility"""
        self.setup_apis()
        self.setup_pinecone()
        self.setup_embedding_model()
        logger.info(f"✅ System initialized on {platform.system()} {platform.release()}")
        
    def setup_apis(self):
        """Setup all API connections"""
        # Gemini API setup
        self.gemini_api_key = os.getenv('GEMINI_API_KEY')
        if not self.gemini_api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables!")
        
        genai.configure(api_key=self.gemini_api_key)
        self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
        logger.info("✅ Gemini API configured successfully")
        
        # Pinecone API setup
        self.pinecone_api_key = os.getenv('PINECONE_API_KEY')
        if not self.pinecone_api_key:
            raise ValueError("PINECONE_API_KEY not found in environment variables!")
        
        logger.info("✅ API keys loaded successfully")
    
    def setup_pinecone(self):
        """Initialize Pinecone vector database"""
        try:
            self.pc = Pinecone(api_key=self.pinecone_api_key)
            self.index_name = "hackrx-document-intelligence"
            self.dimension = 384  # all-MiniLM-L6-v2 embedding dimension
            
            # Check if index exists, create if not
            existing_indexes = [index.name for index in self.pc.list_indexes()]
            
            if self.index_name not in existing_indexes:
                logger.info(f"Creating new Pinecone index: {self.index_name}")
                self.pc.create_index(
                    name=self.index_name,
                    dimension=self.dimension,
                    metric="cosine",
                    spec={
                        "serverless": {
                            "cloud": "aws",
                            "region": "us-east-1"
                        }
                    }
                )
                time.sleep(60)  # Wait for index to be ready
            
            self.index = self.pc.Index(self.index_name)
            logger.info(f"✅ Connected to Pinecone index: {self.index_name}")
            
        except Exception as e:
            logger.error(f"❌ Pinecone setup failed: {e}")
            raise
    
    def setup_embedding_model(self):
        """Initialize embedding model for semantic search"""
        try:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Embedding model loaded successfully")
        except Exception as e:
            logger.error(f"❌ Embedding model setup failed: {e}")
            raise
    
    def download_file(self, url: str, output_path: str) -> bool:
        """Cross-platform file download with retries"""
        try:
            for attempt in range(3):
                try:
                    response = requests.get(url, timeout=30, stream=True)
                    response.raise_for_status()
                    
                    # Ensure directory exists
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    
                    with open(output_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            f.write(chunk)
                    
                    logger.info(f"📥 Downloaded file to {output_path}")
                    return True
                except requests.RequestException as e:
                    logger.warning(f"⚠️ Download attempt {attempt + 1} failed: {e}")
                    if attempt < 2:
                        time.sleep(2)
            
            raise Exception("Download failed after 3 retries")
        except Exception as e:
            logger.error(f"❌ Download error: {e}")
            raise

    def detect_file_type(self, file_path: str) -> str:
        """Cross-platform file type detection"""
        try:
            extension = os.path.splitext(file_path)[1].lower()
            
            # Try reading file header for better detection
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            # PDF detection
            if header.startswith(b'%PDF') or extension == '.pdf':
                return 'PDF'
            # DOCX detection  
            elif header.startswith(b'PK') or extension in ['.docx', '.doc']:
                return 'DOCX'
            # Email detection
            elif extension in ['.eml', '.msg'] or b'From:' in header[:100]:
                return 'EMAIL'
            else:
                # Default based on extension
                if extension == '.pdf':
                    return 'PDF'
                elif extension in ['.docx', '.doc']:
                    return 'DOCX'
                elif extension in ['.eml', '.msg']:
                    return 'EMAIL'
                else:
                    logger.warning(f"⚠️ Unknown file type for {file_path}, assuming PDF")
                    return 'PDF'
                    
        except Exception as e:
            logger.error(f"❌ File type detection error: {e}")
            return 'PDF'  # Default fallback

    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                        continue
            
            logger.info(f"📄 Extracted {len(text)} chars from PDF ({len(pdf_reader.pages)} pages)")
            return text
        except Exception as e:
            logger.error(f"❌ PDF extraction error: {e}")
            return ""

    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += "\t".join(row_text) + "\n"
            
            logger.info(f"📄 Extracted {len(text)} chars from DOCX")
            return text
        except Exception as e:
            logger.error(f"❌ DOCX extraction error: {e}")
            return ""

    def extract_email_text(self, file_path: str) -> str:
        """Extract text from email files"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f)
                
                for part in msg.walk():
                    content_type = part.get_content_type().lower()
                    charset = part.get_content_charset() or 'utf-8'
                    
                    try:
                        if content_type == 'text/plain':
                            payload = part.get_payload(decode=True)
                            part_text = payload.decode(charset, errors='ignore')
                            text += part_text + "\n"
                        elif content_type == 'text/html':
                            payload = part.get_payload(decode=True)
                            html = payload.decode(charset, errors='ignore')
                            soup = BeautifulSoup(html, 'html.parser')
                            part_text = soup.get_text(separator=' ', strip=True)
                            text += part_text + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Email part extraction failed: {e}")
            
            logger.info(f"📧 Extracted {len(text)} chars from email")
            return text
        except Exception as e:
            logger.error(f"❌ Email extraction error: {e}")
            return ""

    def smart_text_chunker(self, text: str, max_chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Intelligent text chunking with semantic awareness"""
        try:
            # Clean and normalize text
            text = re.sub(r'\s+', ' ', text.strip())
            
            if len(text) <= max_chunk_size:
                return [text]
            
            chunks = []
            sentences = re.split(r'(?<=[.!?])\s+', text)
            current_chunk = ""
            
            for sentence in sentences:
                # If adding this sentence exceeds max size
                if len(current_chunk) + len(sentence) > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        
                    # Start new chunk with overlap from previous
                    if chunks and overlap > 0:
                        prev_words = chunks[-1].split()[-overlap//10:]
                        current_chunk = " ".join(prev_words) + " " + sentence
                    else:
                        current_chunk = sentence
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
            
            # Add the last chunk
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            logger.info(f"🔄 Created {len(chunks)} semantic chunks")
            return chunks
        except Exception as e:
            logger.error(f"❌ Text chunking error: {e}")
            return [text]  # Return original text as fallback

    def process_document(self, document_url: str) -> str:
        """Process document and return namespace"""
        try:
            # Generate unique filename and namespace
            filename = document_url.split('/')[-1].split('?')[0]
            if not filename or '.' not in filename:
                filename = f"document_{int(time.time())}.pdf"
            
            output_path = os.path.join("./downloads", filename)
            namespace = hashlib.md5(document_url.encode()).hexdigest()[:16]
            
            # Download file
            self.download_file(document_url, output_path)
            
            # Detect file type and extract text
            file_type = self.detect_file_type(output_path)
            
            if file_type == 'PDF':
                text = self.extract_pdf_text(output_path)
            elif file_type == 'DOCX':
                text = self.extract_docx_text(output_path)
            elif file_type == 'EMAIL':
                text = self.extract_email_text(output_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text.strip():
                raise ValueError("No text extracted from document")
            
            # Create semantic chunks
            chunks = self.smart_text_chunker(text)
            
            # Generate embeddings and store in Pinecone
            self.store_chunks_in_pinecone(chunks, namespace, document_url)
            
            logger.info(f"✅ Document processed: {len(chunks)} chunks in namespace {namespace}")
            return namespace
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            raise

    def store_chunks_in_pinecone(self, chunks: List[str], namespace: str, document_url: str):
        """Store text chunks in Pinecone with embeddings"""
        try:
            # Clear existing namespace
            try:
                self.index.delete(delete_all=True, namespace=namespace)
                time.sleep(2)
            except:
                pass  # Namespace might not exist
            
            # Generate embeddings
            embeddings = self.embedding_model.encode(chunks, convert_to_tensor=False).tolist()
            
            # Prepare vectors for Pinecone
            vectors = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vectors.append({
                    "id": f"{namespace}_{i}",
                    "values": embedding,
                    "metadata": {
                        "text": chunk[:1000],  # Truncate for Pinecone limits
                        "chunk_index": i,
                        "document_url": document_url,
                        "namespace": namespace
                    }
                })
            
            # Upsert in batches
            batch_size = 100
            for i in range(0, len(vectors), batch_size):
                batch = vectors[i:i + batch_size]
                self.index.upsert(vectors=batch, namespace=namespace)
                time.sleep(1)  # Rate limiting
            
            logger.info(f"📤 Stored {len(vectors)} vectors in namespace {namespace}")
            
        except Exception as e:
            logger.error(f"❌ Pinecone storage failed: {e}")
            raise

    def semantic_search(self, query: str, namespace: str, top_k: int = 5) -> List[str]:
        """Perform semantic search and return relevant text chunks"""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query], convert_to_tensor=False).tolist()[0]
            
            # Search in Pinecone
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                namespace=namespace,
                include_metadata=True
            )
            
            # Extract relevant text chunks
            relevant_chunks = []
            for match in results.matches:
                if match.score > 0.3:  # Lowered relevance threshold for better matches
                    relevant_chunks.append(match.metadata.get("text", ""))
            
            logger.info(f"🔍 Found {len(relevant_chunks)} relevant chunks for query")
            return relevant_chunks
            
        except Exception as e:
            logger.error(f"❌ Semantic search failed: {e}")
            return []

    def generate_answer(self, question: str, context_chunks: List[str]) -> str:
        """Generate answer using Gemini with context"""
        try:
            # Prepare context
            context = "\n\n".join(context_chunks[:3])  # Use top 3 most relevant chunks
            
            # Create optimized prompt for token efficiency
            prompt = f"""Based on the following policy document context, answer the question accurately and concisely.

CONTEXT:
{context}

QUESTION: {question}

INSTRUCTIONS:
- Answer based only on the provided context
- Be precise and quote specific terms when mentioned
- If information is not available, state clearly
- Keep response concise but complete

ANSWER:"""

            # Generate response
            response = self.gemini_model.generate_content(prompt)
            answer = response.text.strip()
            
            logger.info(f"🤖 Generated answer for: {question[:50]}...")
            return answer
            
        except Exception as e:
            logger.error(f"❌ Answer generation failed: {e}")
            return "Unable to generate answer due to technical error."

    def process_query(self, document_url: str, questions: List[str]) -> List[str]:
        """Main processing pipeline - matches HackRx API specification"""
        try:
            logger.info(f"🚀 Processing {len(questions)} questions for document")
            
            # Step 1: Process document and get namespace
            namespace = self.process_document(document_url)
            
            # Step 2: Answer each question
            answers = []
            for i, question in enumerate(questions, 1):
                logger.info(f"❓ Processing question {i}/{len(questions)}")
                
                # Semantic search for relevant context
                relevant_chunks = self.semantic_search(question, namespace)
                
                # Generate answer
                answer = self.generate_answer(question, relevant_chunks)
                answers.append(answer)
                
                logger.info(f"✅ Completed question {i}/{len(questions)}")
            
            logger.info(f"🎉 Successfully processed all {len(questions)} questions")
            return answers
            
        except Exception as e:
            logger.error(f"❌ Query processing failed: {e}")
            raise

def main():
    """Test the system with sample data"""
    try:
        # Initialize system
        system = HackRxDocumentIntelligenceSystem()
        
        # Test with sample document and questions from problem statement
        sample_document = "https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D"
        
        sample_questions = [
            "What is the grace period for premium payment under the National Parivar Mediclaim Plus Policy?",
            "What is the waiting period for pre-existing diseases (PED) to be covered?",
            "Does this policy cover maternity expenses, and what are the conditions?"
        ]
        
        # Process queries
        answers = system.process_query(sample_document, sample_questions)
        
        # Display results in required format
        result = {
            "answers": answers
        }
        
        print("\n" + "="*80)
        print("🏆 HACKRX DOCUMENT INTELLIGENCE SYSTEM - API RESPONSE")
        print("="*80)
        print(json.dumps(result, indent=2))
        print("="*80)
        
    except Exception as e:
        logger.error(f"❌ System test failed: {e}")
        raise

if __name__ == "__main__":
    main()
