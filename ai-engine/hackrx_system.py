"""
HackRx Document Intelligence System - Complete Cross-Platform Implementation
Matches exact problem statement requirements and API specification
"""

import requests
import PyPDF2
from docx import Document
import email
from bs4 import BeautifulSoup
import os
import logging
from typing import Tuple, List, Dict, Any, Optional
import re
import time
import hashlib
import json
from dotenv import load_dotenv
import google.generativeai as genai
import pinecone
from sentence_transformers import SentenceTransformer
import platform
import asyncio
import aiohttp
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from functools import lru_cache
import numpy as np

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class HackRxDocumentIntelligenceSystem:
    """Complete Document Intelligence System matching HackRx specifications"""
    
    def __init__(self):
        """Initialize the system with cross-platform compatibility"""
        self.setup_apis()
        self.setup_pinecone()
        self.setup_embedding_model()
        # Cache for processed documents to avoid reprocessing
        self.document_cache = {}
        self.namespace_cache = {}
        # Store document text for fallback search
        self.document_text_cache = {}
        self.document_chunks_cache = {}
        # Thread pool for parallel processing
        self.executor = ThreadPoolExecutor(max_workers=4)
        logger.info(f"✅ System initialized on {platform.system()} {platform.release()}")
        logger.info("🚀 Performance optimizations enabled!")
        
    def setup_apis(self):
        """Setup all API connections"""
        # Gemini API setup
        self.gemini_api_key = os.getenv('GEMINI_API_KEY')
        if not self.gemini_api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables!")
        
        genai.configure(api_key=self.gemini_api_key)
        self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
        logger.info("✅ Gemini API configured successfully")
        
        # Pinecone API setup
        self.pinecone_api_key = os.getenv('PINECONE_API_KEY')
        if not self.pinecone_api_key:
            raise ValueError("PINECONE_API_KEY not found in environment variables!")
        
        logger.info("✅ API keys loaded successfully")
    
    def setup_pinecone(self):
        """Initialize vector storage with local fallback for speed"""
        try:
            logger.info("🔄 Setting up vector storage...")
            # Use local storage for faster performance and avoid Pinecone issues
            self.use_pinecone = False
            self.local_embeddings = {}  # Store embeddings locally
            self.document_texts = {}    # Store full document texts
            self.index_name = "local-storage"
            logger.info("✅ Using local vector storage (faster and more reliable)")
            
        except Exception as e:
            logger.error(f"❌ Vector storage setup failed: {e}")
            raise
    
    def setup_embedding_model(self):
        """Initialize embedding model for semantic search"""
        try:
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Embedding model loaded successfully")
        except Exception as e:
            logger.error(f"❌ Embedding model setup failed: {e}")
            raise
    
    def download_file(self, url: str, output_path: str) -> bool:
        """Optimized cross-platform file download with caching"""
        try:
            # Check if file already exists and is valid
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info(f"📁 Using cached file: {output_path}")
                return True
                
            for attempt in range(2):  # Reduced from 3 attempts
                try:
                    response = requests.get(url, timeout=15, stream=True)  # Reduced timeout
                    response.raise_for_status()
                    
                    # Ensure directory exists
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    
                    with open(output_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=16384):  # Larger chunks
                            f.write(chunk)
                    
                    logger.info(f"📥 Downloaded file to {output_path}")
                    return True
                except requests.RequestException as e:
                    logger.warning(f"⚠️ Download attempt {attempt + 1} failed: {e}")
                    if attempt < 1:
                        time.sleep(1)  # Reduced sleep time
            
            raise Exception("Download failed after 2 retries")
        except Exception as e:
            logger.error(f"❌ Download error: {e}")
            raise

    def detect_file_type(self, file_path: str) -> str:
        """Cross-platform file type detection"""
        try:
            extension = os.path.splitext(file_path)[1].lower()
            
            # Try reading file header for better detection
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            # PDF detection
            if header.startswith(b'%PDF') or extension == '.pdf':
                return 'PDF'
            # DOCX detection  
            elif header.startswith(b'PK') or extension in ['.docx', '.doc']:
                return 'DOCX'
            # Email detection
            elif extension in ['.eml', '.msg'] or b'From:' in header[:100]:
                return 'EMAIL'
            else:
                # Default based on extension
                if extension == '.pdf':
                    return 'PDF'
                elif extension in ['.docx', '.doc']:
                    return 'DOCX'
                elif extension in ['.eml', '.msg']:
                    return 'EMAIL'
                else:
                    logger.warning(f"⚠️ Unknown file type for {file_path}, assuming PDF")
                    return 'PDF'
                    
        except Exception as e:
            logger.error(f"❌ File type detection error: {e}")
            return 'PDF'  # Default fallback

    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract page {page_num}: {e}")
                        continue
            
            logger.info(f"📄 Extracted {len(text)} chars from PDF ({len(pdf_reader.pages)} pages)")
            return text
        except Exception as e:
            logger.error(f"❌ PDF extraction error: {e}")
            return ""

    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += "\t".join(row_text) + "\n"
            
            logger.info(f"📄 Extracted {len(text)} chars from DOCX")
            return text
        except Exception as e:
            logger.error(f"❌ DOCX extraction error: {e}")
            return ""

    def extract_email_text(self, file_path: str) -> str:
        """Extract text from email files"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f)
                
                for part in msg.walk():
                    content_type = part.get_content_type().lower()
                    charset = part.get_content_charset() or 'utf-8'
                    
                    try:
                        if content_type == 'text/plain':
                            payload = part.get_payload(decode=True)
                            part_text = payload.decode(charset, errors='ignore')
                            text += part_text + "\n"
                        elif content_type == 'text/html':
                            payload = part.get_payload(decode=True)
                            html = payload.decode(charset, errors='ignore')
                            soup = BeautifulSoup(html, 'html.parser')
                            part_text = soup.get_text(separator=' ', strip=True)
                            text += part_text + "\n"
                    except Exception as e:
                        logger.warning(f"⚠️ Email part extraction failed: {e}")
            
            logger.info(f"📧 Extracted {len(text)} chars from email")
            return text
        except Exception as e:
            logger.error(f"❌ Email extraction error: {e}")
            return ""

    def smart_text_chunker(self, text: str, max_chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Optimized text chunking with semantic awareness"""
        try:
            # Clean and normalize text more efficiently
            text = re.sub(r'\s+', ' ', text.strip())
            
            if len(text) <= max_chunk_size:
                return [text]
            
            chunks = []
            # Use simpler sentence splitting for speed
            sentences = re.split(r'[.!?]+\s+', text)
            current_chunk = ""
            
            for sentence in sentences:
                # If adding this sentence exceeds max size
                if len(current_chunk) + len(sentence) > max_chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                        
                    # Start new chunk with minimal overlap for speed
                    if chunks and overlap > 0:
                        prev_words = chunks[-1].split()[-5:]  # Fixed overlap
                        current_chunk = " ".join(prev_words) + " " + sentence
                    else:
                        current_chunk = sentence
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
            
            # Add the last chunk
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            logger.info(f"🔄 Created {len(chunks)} semantic chunks")
            return chunks
        except Exception as e:
            logger.error(f"❌ Text chunking error: {e}")
            return [text]  # Return original text as fallback

    def process_document(self, document_url: str) -> str:
        """Optimized document processing with caching"""
        try:
            # Check cache first
            url_hash = hashlib.md5(document_url.encode()).hexdigest()
            if url_hash in self.document_cache:
                logger.info(f"📁 Using cached document processing for {document_url[:50]}...")
                return self.document_cache[url_hash]
            
            # Generate unique filename and namespace
            filename = document_url.split('/')[-1].split('?')[0]
            if not filename or '.' not in filename:
                filename = f"document_{int(time.time())}.pdf"
            
            output_path = os.path.join("./downloads", filename)
            namespace = url_hash[:16]
            
            # Download file (with caching)
            self.download_file(document_url, output_path)
            
            # Detect file type and extract text
            file_type = self.detect_file_type(output_path)
            
            if file_type == 'PDF':
                text = self.extract_pdf_text(output_path)
            elif file_type == 'DOCX':
                text = self.extract_docx_text(output_path)
            elif file_type == 'EMAIL':
                text = self.extract_email_text(output_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text.strip():
                raise ValueError("No text extracted from document")
            
            # Create semantic chunks
            chunks = self.smart_text_chunker(text)
            
            # Cache the document text and chunks for fallback search
            self.document_text_cache[url_hash] = text
            self.document_chunks_cache[url_hash] = chunks
            
            # Generate embeddings and store in Pinecone
            self.store_chunks_in_pinecone(chunks, namespace, document_url)
            
            # Cache the result
            self.document_cache[url_hash] = namespace
            
            logger.info(f"✅ Document processed: {len(chunks)} chunks in namespace {namespace}")
            return namespace
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            raise

    def store_chunks_in_pinecone(self, chunks: List[str], namespace: str, document_url: str):
        """Store text chunks locally for faster performance"""
        try:
            # Check if namespace already exists
            if namespace in self.namespace_cache:
                logger.info(f"📁 Namespace {namespace} already exists, skipping storage")
                return
                
            logger.info(f"🔄 Storing {len(chunks)} chunks locally...")
            
            # Store full text for keyword search
            self.document_texts[namespace] = " ".join(chunks)
            
            # Generate embeddings for semantic search
            embeddings = self.embedding_model.encode(
                chunks, 
                convert_to_tensor=False, 
                show_progress_bar=False,
                batch_size=32
            ).tolist()
            
            # Store locally
            self.local_embeddings[namespace] = {
                'chunks': chunks,
                'embeddings': embeddings,
                'document_url': document_url
            }
            
            # Cache namespace
            self.namespace_cache[namespace] = True
            
            logger.info(f"✅ Stored {len(chunks)} chunks locally in namespace {namespace}")
            
        except Exception as e:
            logger.error(f"❌ Local storage failed: {e}")
            # Don't raise - continue without vector search

    def semantic_search(self, query: str, namespace: str, top_k: int = 5) -> List[str]:
        """Perform local semantic search for fast results"""
        try:
            # Use local embeddings for semantic search
            if namespace in self.local_embeddings:
                data = self.local_embeddings[namespace]
                chunks = data['chunks']
                stored_embeddings = data['embeddings']
                
                # Generate query embedding
                query_embedding = self.embedding_model.encode([query], convert_to_tensor=False).tolist()[0]
                
                # Calculate similarities
                import numpy as np
                similarities = []
                for i, chunk_embedding in enumerate(stored_embeddings):
                    similarity = np.dot(query_embedding, chunk_embedding) / (
                        np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
                    )
                    similarities.append((similarity, i))
                
                # Sort by similarity and get top results
                similarities.sort(reverse=True)
                relevant_chunks = []
                
                for similarity, idx in similarities[:top_k]:
                    if similarity > 0.3:  # Relevance threshold
                        relevant_chunks.append(chunks[idx])
                
                logger.info(f"🔍 Found {len(relevant_chunks)} relevant chunks locally")
                return relevant_chunks
            
            # Fallback to text search
            return self._fallback_text_search(query, namespace)
            
        except Exception as e:
            logger.error(f"❌ Semantic search failed: {e}")
            return self._fallback_text_search(query, namespace)
    
    def _fallback_text_search(self, query: str, namespace: str) -> List[str]:
        """Enhanced fallback text search using local storage"""
        try:
            # Check local embeddings first
            if namespace in self.local_embeddings:
                chunks = self.local_embeddings[namespace]['chunks']
                query_lower = query.lower()
                
                # Simple keyword matching - find chunks containing query terms
                relevant_chunks = []
                query_words = query_lower.split()
                
                for chunk in chunks:
                    chunk_lower = chunk.lower()
                    score = sum(1 for word in query_words if word in chunk_lower)
                    if score > 0:
                        relevant_chunks.append((chunk, score))
                
                # Sort by relevance score and return top chunks
                relevant_chunks.sort(key=lambda x: x[1], reverse=True)
                result = [chunk for chunk, score in relevant_chunks[:3]]
                
                logger.info(f"🔍 Fallback search found {len(result)} relevant chunks")
                return result if result else chunks[:3]  # Return first 3 chunks if no matches
            
            # Check full document text
            if namespace in self.document_texts:
                full_text = self.document_texts[namespace]
                # Split into manageable chunks and search
                chunks = [full_text[i:i+1000] for i in range(0, len(full_text), 800)]
                query_lower = query.lower()
                
                relevant_chunks = []
                for chunk in chunks:
                    if any(word in chunk.lower() for word in query_lower.split()):
                        relevant_chunks.append(chunk)
                
                return relevant_chunks[:3] if relevant_chunks else chunks[:3]
            
            logger.warning("⚠️ No cached content found for fallback search")
            return ["Document content temporarily unavailable for semantic search."]
            
        except Exception as e:
            logger.error(f"❌ Fallback search failed: {e}")
            return ["Unable to perform text search due to technical error."]

    def generate_answer(self, question: str, context_chunks: List[str]) -> str:
        """Generate answer using Gemini with professional insurance company context"""
        try:
            # Prepare context
            context = "\n\n".join(context_chunks[:3])  # Use top 3 most relevant chunks
            
            # Create professional insurance company prompt
            prompt = f"""You are a highly capable, polite, and trustworthy virtual assistant for a leading multinational insurance company.

Your role is to assist customers by answering their questions based on internal documentation (shown as "Context" below). Always respond in a clear, professional, and supportive tone, maintaining the company's reputation for excellence and care.

Your response should:
- Accurately address the customer's question using only the information from the context
- Be concise, structured, and easy to understand
- Include bullet points or paragraphs as appropriate
- Avoid guessing or fabricating information
- End with a courteous support message

Context:
{context}

Customer Question:
{question}

Final Output Format:
[Your structured answer here]

Please close your message with:
"If you need further assistance, feel free to contact our customer support team at support@example.com or call us at +1-800-123-4567. We're here to help 24/7."
"""

            # Generate response
            response = self.gemini_model.generate_content(prompt)
            answer = response.text.strip()
            
            logger.info(f"🤖 Generated answer for: {question[:50]}...")
            return answer
            
        except Exception as e:
            logger.error(f"❌ Answer generation failed: {e}")
            return "Unable to generate answer due to technical error. If you need further assistance, feel free to contact our customer support team at support@example.com or call us at +1-800-123-4567. We're here to help 24/7."

    def process_query(self, document_url: str, questions: List[str]) -> List[str]:
        """Optimized main processing pipeline with parallel question processing"""
        try:
            logger.info(f"🚀 Processing {len(questions)} questions for document")
            
            # Step 1: Process document once (with caching)
            namespace = self.process_document(document_url)
            
            # Step 2: Process questions in parallel for speed
            logger.info(f"🔄 Processing questions in parallel...")
            
            def process_single_question(question_data):
                i, question = question_data
                logger.info(f"❓ Processing question {i}/{len(questions)}: {question[:50]}...")
                
                # Semantic search for relevant context
                relevant_chunks = self.semantic_search(question, namespace)
                
                # Generate answer
                answer = self.generate_answer(question, relevant_chunks)
                
                logger.info(f"✅ Completed question {i}/{len(questions)}")
                return answer
            
            # Use ThreadPoolExecutor for parallel processing
            with ThreadPoolExecutor(max_workers=3) as executor:
                question_data = [(i+1, q) for i, q in enumerate(questions)]
                answers = list(executor.map(process_single_question, question_data))
            
            logger.info(f"🎉 Successfully processed all {len(questions)} questions in parallel")
            return answers
            
        except Exception as e:
            logger.error(f"❌ Query processing failed: {e}")
            raise

def main():
    """Test the system with sample data"""
    try:
        # Initialize system
        system = HackRxDocumentIntelligenceSystem()
        
        # Test with sample document and questions from problem statement
        sample_document = "https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D"
        
        sample_questions = [
            "What is the grace period for premium payment under the National Parivar Mediclaim Plus Policy?",
            "What is the waiting period for pre-existing diseases (PED) to be covered?",
            "Does this policy cover maternity expenses, and what are the conditions?"
        ]
        
        # Process queries
        answers = system.process_query(sample_document, sample_questions)
        
        # Display results in required format
        result = {
            "answers": answers
        }
        
        print("\n" + "="*80)
        print("🏆 HACKRX DOCUMENT INTELLIGENCE SYSTEM - API RESPONSE")
        print("="*80)
        print(json.dumps(result, indent=2))
        print("="*80)
        
    except Exception as e:
        logger.error(f"❌ System test failed: {e}")
        raise

if __name__ == "__main__":
    main()
